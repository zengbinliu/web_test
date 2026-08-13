#!/usr/bin/env python3
"""Parse PRD documents (md/txt/docx/pdf) into structured JSON for gen_tc."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

MODULE_PATTERNS = [
    re.compile(r"^#{1,4}\s+(.+)"),
    re.compile(r"^第[一二三四五六七八九十\d]+[章节条]\s*(.+)"),
    re.compile(r"^§\s*[\d.]+\s*(.+)"),
    re.compile(r"^\d+(?:\.\d+)*\s+(.{4,80})$"),
]

RULE_KEYWORDS = ("必须", "应当", "不得", "禁止", "仅允许", "如果", "当", "规则", "逻辑", "应")
BOUNDARY_KEYWORDS = ("边界", "范围", "最大", "最小", "不超过", "至少", "上限", "下限", "≤", "≥", "最多", "最少")
EXCEPTION_KEYWORDS = ("异常", "错误", "失败", "提示", "报错", "不可用", "拒绝")
QUESTION_PATTERNS = (
    re.compile(r"[？?]\s*$"),
    re.compile(r"^TODO\b", re.I),
    re.compile(r"待确认|待产品|TBD|待定", re.I),
)


def _dedupe(items: list[str], limit: int) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for item in items:
        key = item[:120]
        if key in seen:
            continue
        seen.add(key)
        out.append(item)
        if len(out) >= limit:
            break
    return out


def parse_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("Missing dependency: pip install python-docx") from exc

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_pdf(path: Path) -> tuple[str, list[str]]:
    warnings: list[str] = []
    text = ""

    try:
        import pdfplumber

        parts: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                chunk = page.extract_text() or ""
                if chunk.strip():
                    parts.append(chunk)
        text = "\n".join(parts)
    except ImportError:
        warnings.append("pdfplumber not installed")
    except Exception as exc:
        warnings.append(f"pdfplumber failed: {exc}")

    if not text.strip():
        try:
            import fitz

            parts = []
            with fitz.open(path) as doc:
                for page in doc:
                    chunk = page.get_text() or ""
                    if chunk.strip():
                        parts.append(chunk)
            text = "\n".join(parts)
        except ImportError:
            warnings.append("pymupdf not installed; pip install pymupdf")
        except Exception as exc:
            warnings.append(f"pymupdf failed: {exc}")

    if not text.strip():
        raise SystemExit(
            "Could not extract text from PDF. Install pdfplumber or pymupdf, "
            "or provide a text/md copy of the requirement."
        )
    return text, warnings


def extract_structure(text: str) -> dict[str, list[str]]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    modules: list[str] = []
    rules: list[str] = []
    boundaries: list[str] = []
    exceptions: list[str] = []
    open_questions: list[str] = []

    for line in lines:
        for pattern in MODULE_PATTERNS:
            match = pattern.match(line)
            if match and len(match.group(1)) < 80:
                modules.append(match.group(1).strip())
                break

        if len(line) < 300:
            if any(keyword in line for keyword in RULE_KEYWORDS):
                rules.append(line)
            if any(keyword in line for keyword in BOUNDARY_KEYWORDS):
                boundaries.append(line)
            if any(keyword in line for keyword in EXCEPTION_KEYWORDS):
                exceptions.append(line)

        if any(pattern.search(line) for pattern in QUESTION_PATTERNS):
            open_questions.append(line)

    return {
        "modules": _dedupe(modules, 30),
        "rules": _dedupe(rules, 40),
        "boundaries": _dedupe(boundaries, 30),
        "exceptions": _dedupe(exceptions, 30),
        "open_questions": _dedupe(open_questions, 20),
    }


def parse_prd(path: Path, *, include_text: bool = True) -> dict[str, Any]:
    suffix = path.suffix.lower()
    warnings: list[str] = []

    if suffix in (".md", ".txt"):
        text = parse_text_file(path)
    elif suffix == ".docx":
        text = parse_docx(path)
    elif suffix == ".pdf":
        text, warnings = parse_pdf(path)
    else:
        raise SystemExit(f"Unsupported format: {suffix}. Use .md .txt .docx .pdf")

    structure = extract_structure(text)
    if len(text) < 100:
        warnings.append("Extracted text is very short; document may be incomplete or scanned")

    result: dict[str, Any] = {
        "source": str(path.resolve()),
        "format": suffix.lstrip("."),
        "char_count": len(text),
        "line_count": len(text.splitlines()),
        "parse_warnings": warnings,
        "structure": structure,
    }
    if include_text:
        result["text"] = text
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description="Parse PRD for gen_tc")
    parser.add_argument("--input", "-i", required=True, help="PRD file path")
    parser.add_argument("--output", "-o", help="Output JSON path; default stdout")
    parser.add_argument(
        "--no-text",
        action="store_true",
        help="Omit full text from JSON (structure summary only)",
    )
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.is_file():
        print(f"Input not found: {input_path}", file=sys.stderr)
        return 1

    result = parse_prd(input_path, include_text=not args.no_text)
    payload = json.dumps(result, ensure_ascii=False, indent=2)

    if args.output:
        output_path = Path(args.output).expanduser().resolve()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(payload + "\n", encoding="utf-8")
        print(f"Wrote structure -> {output_path}")
    else:
        print(payload)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
